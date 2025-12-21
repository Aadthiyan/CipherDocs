"""
Text extraction logic for various file formats (PDF, DOCX, TXT).
"""

import io
import logging
from typing import BinaryIO, Dict, Any, Optional, Tuple
import chardet

# Pydantic is used for structured logging and potential API integrations
from pydantic import BaseModel

logger = logging.getLogger(__name__)

class ExtractedContent(BaseModel):
    text: str
    metadata: Dict[str, Any] = {}
    page_count: int = 0
    error: Optional[str] = None

class TextExtractor:
    """Base class and factory for text extraction"""
    
    @staticmethod
    def extract(file_data: BinaryIO, doc_type: str) -> ExtractedContent:
        """
        Extract text from file content based on document type.
        
        Args:
            file_data: Binary file content
            doc_type: Document extension/type (pdf, docx, txt)
            
        Returns:
            ExtractedContent object with text and metadata
        """
        file_data.seek(0)
        doc_type = doc_type.lower()
        
        try:
            if doc_type == 'pdf':
                return TextExtractor._extract_pdf(file_data)
            elif doc_type in ['docx', 'doc']:
                return TextExtractor._extract_docx(file_data)
            elif doc_type == 'txt':
                return TextExtractor._extract_txt(file_data)
            else:
                raise ValueError(f"Unsupported document type: {doc_type}")
        except Exception as e:
            logger.error(f"Extraction failed for type {doc_type}: {str(e)}")
            return ExtractedContent(
                text="",
                error=str(e)
            )

    @staticmethod
    def _extract_pdf(file_data: BinaryIO) -> ExtractedContent:
        """Extract text from PDF using pypdf/pdfplumber"""
        import pypdf
        
        text_content = []
        metadata = {}
        page_count = 0
        
        try:
            reader = pypdf.PdfReader(file_data)
            page_count = len(reader.pages)
            metadata = reader.metadata if reader.metadata else {}
            
            # Clean metadata keys
            if metadata:
                metadata = {k.strip('/'): v for k, v in metadata.items()}
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    text_content.append(text)
                else:
                    # Fallback or notation for empty pages (could be images)
                    logger.warning(f"Empty text on page {i+1} of PDF")
            
            return ExtractedContent(
                text="\n\n".join(text_content),
                metadata=metadata,
                page_count=page_count
            )
            
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise e

    @staticmethod
    def _extract_docx(file_data: BinaryIO) -> ExtractedContent:
        """Extract text from DOCX using python-docx"""
        import docx
        
        try:
            doc = docx.Document(file_data)
            text_content = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    text_content.append(" | ".join(row_data))
            
            # Extract core properties as metadata
            metadata = {}
            if doc.core_properties:
                props = doc.core_properties
                if props.author: metadata['author'] = props.author
                if props.title: metadata['title'] = props.title
                if props.created: metadata['created'] = str(props.created)
                
            return ExtractedContent(
                text="\n\n".join(text_content),
                metadata=metadata,
                page_count=1 # Approximate or not applicable
            )
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise e

    @staticmethod
    def _extract_txt(file_data: BinaryIO) -> ExtractedContent:
        """Extract text from TXT with encoding detection"""
        raw_data = file_data.read()
        
        # Detect encoding
        result = chardet.detect(raw_data)
        encoding = result['encoding'] or 'utf-8'
        
        try:
            text = raw_data.decode(encoding)
            return ExtractedContent(
                text=text,
                metadata={'encoding': encoding},
                page_count=1
            )
        except UnicodeDecodeError:
            # Fallback to utf-8 with replacement
            text = raw_data.decode('utf-8', errors='replace')
            return ExtractedContent(
                text=text,
                metadata={'encoding': 'utf-8-replace'},
                page_count=1
            )
